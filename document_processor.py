import os
import re
import email
from email import policy
from typing import List, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document
import tiktoken

from config import Config

class DocumentProcessor:
    """Handles processing of various document types into chunks"""
    
    def __init__(self):
        self.config = Config()
        self.encoding = tiktoken.get_encoding("cl100k_base")
        
        # Performance optimizations - compile regex patterns once
        self._compiled_patterns = {
            'list_item': re.compile(r'^\s*[•\-\*]|^\s*\d+[.):]'),
            'monetary': re.compile(r'₹|rupees?|\d+\s*lakh|\d+\s*crore'),
            'percentage': re.compile(r'\d+\s*%|\d+\s*(months?|years?|days?)'),
            'sentence_split': re.compile(r'(?<=[.!?])\s+')
        }
        
        # Cache for tokenization to avoid repeated encoding
        self._token_cache = {}
        
        # Pre-compiled keyword sets for faster lookup
        self._important_terms_set = {
            'coverage', 'benefit', 'premium', 'claim', 'exclusion', 'waiting period',
            'grace period', 'deductible', 'copayment', 'pre-existing', 'maternity',
            'room rent', 'icu', 'surgery', 'hospitalization'
        }
    
    def process_document(self, file_path: str, filename: str) -> List[Dict[str, Any]]:
        """
        Process a document and return chunks with metadata - OPTIMIZED
        
        Args:
            file_path: Path to the document file
            filename: Original filename
            
        Returns:
            List of document chunks with metadata
        """
        # Clear cache if it gets too large
        if len(self._token_cache) > 1000:
            self.clear_cache()
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif file_extension in ['.docx', '.doc']:
            text = self._extract_docx_text(file_path)
        elif file_extension == '.txt':
            text = self._extract_txt_text(file_path)
        elif file_extension == '.eml':
            text = self._extract_email_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Clean and preprocess text
        text = self._clean_text(text)
        
        # Create chunks
        chunks = self._create_chunks(text, filename)
        
        return chunks
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
        
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting Word document text: {str(e)}")
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error extracting text file: {str(e)}")
    
    def _extract_email_text(self, file_path: str) -> str:
        """Extract text from email file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                msg = email.message_from_file(file)
            
            text = ""
            
            # Extract headers
            text += f"From: {msg.get('From', 'Unknown')}\n"
            text += f"To: {msg.get('To', 'Unknown')}\n"
            text += f"Subject: {msg.get('Subject', 'No Subject')}\n"
            text += f"Date: {msg.get('Date', 'Unknown')}\n\n"
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_payload(decode=True)
                        if body:
                            text += body.decode('utf-8', errors='ignore')
            else:
                body = msg.get_payload(decode=True)
                if body:
                    text += body.decode('utf-8', errors='ignore')
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting email text: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and preprocess text with better normalization for semantic search"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Normalize common variations and abbreviations
        text = self._normalize_text_variations(text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>\|\\]', '', text)
        
        # Remove multiple consecutive newlines
        text = re.sub(r'\n+', '\n', text)
        
        # Convert to lowercase for better semantic matching (preserve original structure)
        # We'll store both original and normalized versions
        return text.strip()
    
    def _normalize_text_variations(self, text: str) -> str:
        """Normalize common text variations for better semantic matching"""
        # Common insurance/medical term normalizations
        normalizations = {
            r'\bpre-existing\b': 'preexisting',
            r'\bpre existing\b': 'preexisting', 
            r'\bco-payment\b': 'copayment',
            r'\bco payment\b': 'copayment',
            r'\bdeductible\b': 'deductible',
            r'\bdeductable\b': 'deductible',
            r'\bbenefits?\b': 'benefit',
            r'\bcoverages?\b': 'coverage',
            r'\bpolicies\b': 'policy',
            r'\bpremiums?\b': 'premium',
            r'\bclaims?\b': 'claim',
            r'\bprocedures?\b': 'procedure',
            r'\btreatments?\b': 'treatment',
            r'\bhospitalizations?\b': 'hospitalization',
            r'\bsurgeries\b': 'surgery',
            r'\bsurgical\b': 'surgery',
            r'\bmedical\b': 'medical',
            r'\bdental\b': 'dental',
            r'\bmaternity\b': 'maternity',
            r'\bwaiting period\b': 'waiting period',
            r'\bgrace period\b': 'grace period',
            r'\bno claim discount\b': 'no claim discount',
            r'\bNCD\b': 'no claim discount',
            r'\bAYUSH\b': 'ayush',
            r'\bICU\b': 'icu intensive care',
            r'\bOPD\b': 'opd outpatient',
            r'\bIPD\b': 'ipd inpatient'
        }
        
        # Apply normalizations (case insensitive)
        for pattern, replacement in normalizations.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text
    
    def _create_chunks(self, text: str, source: str) -> List[Dict[str, Any]]:
        """Create semantic chunks with edge windows for better query handling"""
        try:
            # First, analyze document structure
            semantic_blocks = self._analyze_document_structure(text)
            
            # Create semantic chunks with edge windows
            chunks = self._create_semantic_chunks_with_edges(semantic_blocks, source)
            
            # Fallback to simple chunking if semantic chunking produces too few chunks
            if len(chunks) < 3 and len(text) > 3000:  # Should have more chunks for large docs
                print(f"⚠️ Semantic chunking produced only {len(chunks)} chunks, falling back to simple chunking")
                chunks = self._create_simple_chunks(text, source)
            
            return chunks
        except Exception as e:
            print(f"❌ Semantic chunking failed: {e}, falling back to simple chunking")
            return self._create_simple_chunks(text, source)
    
    def _analyze_document_structure(self, text: str) -> List[Dict[str, Any]]:
        """Analyze document structure to identify semantic blocks - OPTIMIZED"""
        blocks = []
        
        # Split text into paragraphs first
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        # Batch process paragraphs for better performance
        paragraph_data = self._batch_process_paragraphs(paragraphs)
        
        for i, (paragraph, paragraph_lower, token_count) in enumerate(paragraph_data):
            block_type = self._identify_block_type_optimized(paragraph, paragraph_lower)
            importance_score = self._calculate_importance_score_optimized(paragraph, paragraph_lower)
            
            # Only split into sentences if needed (for very long paragraphs)
            sentences = self._split_into_sentences_cached(paragraph) if len(paragraph) > 500 else [paragraph]
            
            block = {
                'content': paragraph,
                'sentences': sentences,
                'block_type': block_type,
                'importance_score': importance_score,
                'position': i,
                'token_count': token_count
            }
            blocks.append(block)
        
        return blocks
    
    def _batch_process_paragraphs(self, paragraphs: List[str]) -> List[tuple]:
        """Batch process paragraphs for tokenization and lowercasing"""
        batch_data = []
        
        for paragraph in paragraphs:
            # Cache tokenization results
            if paragraph in self._token_cache:
                token_count = self._token_cache[paragraph]
            else:
                token_count = len(self.encoding.encode(paragraph))
                # Only cache if reasonable size to avoid memory issues
                if len(self._token_cache) < 1000:
                    self._token_cache[paragraph] = token_count
            
            paragraph_lower = paragraph.lower()
            batch_data.append((paragraph, paragraph_lower, token_count))
        
        return batch_data
    
    def _identify_block_type_optimized(self, text: str, text_lower: str) -> str:
        """Identify the type of content block - OPTIMIZED"""
        text_len = len(text)
        
        # Headers and titles (short, often capitalized)
        if text_len < 100 and (text.isupper() or text.istitle()):
            return 'header'
        
        # Lists (bullet points, numbered) - use compiled regex
        if self._compiled_patterns['list_item'].match(text):
            return 'list_item'
        
        # Tables (contains pipe separators or multiple tabs)
        if '|' in text or text.count('\t') > 3:
            return 'table'
        
        # Policy sections (contains specific keywords) - use set intersection
        policy_keywords = {'coverage', 'benefit', 'premium', 'claim', 'exclusion', 'condition'}
        text_words = set(text_lower.split())
        if policy_keywords & text_words:  # Fast set intersection
            return 'policy_section'
        
        # Definitions (contains "means" or "refers to")
        if 'means' in text_lower or 'refers to' in text_lower or 'defined as' in text_lower:
            return 'definition'
        
        # Default to paragraph
        return 'paragraph'
    
    def _identify_block_type(self, text: str) -> str:
        """Identify the type of content block - LEGACY METHOD"""
        return self._identify_block_type_optimized(text, text.lower().strip())
    
    def _calculate_importance_score_optimized(self, text: str, text_lower: str) -> float:
        """Calculate importance score for content prioritization - OPTIMIZED"""
        score = 1.0
        text_len = len(text)
        
        # Boost score for important keywords using pre-compiled set
        text_words = set(text_lower.split())
        matching_terms = self._important_terms_set & text_words
        score += len(matching_terms) * 0.2
        
        # Boost for monetary amounts using compiled regex
        if self._compiled_patterns['monetary'].search(text_lower):
            score += 0.3
        
        # Boost for percentages and numbers using compiled regex
        if self._compiled_patterns['percentage'].search(text_lower):
            score += 0.2
        
        # Reduce score for very short or very long blocks
        if text_len < 50:
            score *= 0.8
        elif text_len > 1000:
            score *= 0.9
        
        return min(score, 3.0)  # Cap at 3.0
    
    def _calculate_importance_score(self, text: str) -> float:
        """Calculate importance score for content prioritization - LEGACY METHOD"""
        return self._calculate_importance_score_optimized(text, text.lower())
    
    def _create_semantic_chunks_with_edges(self, blocks: List[Dict[str, Any]], source: str) -> List[Dict[str, Any]]:
        """Create semantic chunks with edge windows for better context preservation"""
        chunks = []
        current_chunk_blocks = []
        current_tokens = 0
        chunk_id = 0
        
        for i, block in enumerate(blocks):
            block_tokens = block['token_count']
            
            # Check if adding this block would exceed chunk size
            if current_tokens + block_tokens > self.config.CHUNK_SIZE and current_chunk_blocks:
                # Create chunk with edge windows
                chunk_content = self._build_chunk_with_edges(current_chunk_blocks, blocks, i)
                
                chunk_metadata = self._create_enhanced_chunk_metadata(
                    chunk_content, source, chunk_id, current_chunk_blocks
                )
                chunks.append(chunk_metadata)
                
                # Start new chunk with overlap from previous chunk
                overlap_blocks = self._get_edge_window_blocks(current_chunk_blocks, blocks, i)
                current_chunk_blocks = overlap_blocks + [block]
                current_tokens = sum(b['token_count'] for b in current_chunk_blocks)
                chunk_id += 1
            else:
                current_chunk_blocks.append(block)
                current_tokens += block_tokens
        
        # Handle the last chunk
        if current_chunk_blocks:
            chunk_content = self._build_chunk_with_edges(current_chunk_blocks, blocks, len(blocks))
            chunk_metadata = self._create_enhanced_chunk_metadata(
                chunk_content, source, chunk_id, current_chunk_blocks
            )
            chunks.append(chunk_metadata)
        
        return chunks
    
    def _create_simple_chunks(self, text: str, source: str) -> List[Dict[str, Any]]:
        """Fallback simple chunking method for when semantic chunking fails"""
        chunks = []
        sentences = self._split_into_sentences_cached(text)
        
        current_chunk = []
        current_tokens = 0
        chunk_id = 0
        
        for sentence in sentences:
            sentence_tokens = len(self.encoding.encode(sentence))
            
            # If adding this sentence would exceed chunk size, finalize current chunk
            if current_tokens + sentence_tokens > self.config.CHUNK_SIZE and current_chunk:
                chunk_content = ' '.join(current_chunk)
                
                # Add overlap from previous chunk
                if chunk_id > 0 and len(current_chunk) > 1:
                    overlap_size = min(2, len(current_chunk) // 3)
                    overlap_content = ' '.join(current_chunk[-overlap_size:])
                    chunk_content = f"{overlap_content} {chunk_content}"
                
                chunk_metadata = {
                    'content': chunk_content,
                    'normalized_content': self._normalize_text(chunk_content),
                    'source': source,
                    'chunk_id': chunk_id,
                    'token_count': len(self.encoding.encode(chunk_content)),
                    'block_types': ['paragraph'],
                    'importance_score': 1.0,
                    'position': chunk_id
                }
                chunks.append(chunk_metadata)
                
                # Start new chunk with overlap
                if len(current_chunk) > 1:
                    overlap_size = min(2, len(current_chunk) // 3)
                    current_chunk = current_chunk[-overlap_size:] + [sentence]
                    current_tokens = sum(len(self.encoding.encode(s)) for s in current_chunk)
                else:
                    current_chunk = [sentence]
                    current_tokens = sentence_tokens
                chunk_id += 1
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens
        
        # Handle the last chunk
        if current_chunk:
            chunk_content = ' '.join(current_chunk)
            chunk_metadata = {
                'content': chunk_content,
                'normalized_content': self._normalize_text(chunk_content),
                'source': source,
                'chunk_id': chunk_id,
                'token_count': len(self.encoding.encode(chunk_content)),
                'block_types': ['paragraph'],
                'importance_score': 1.0,
                'position': chunk_id
            }
            chunks.append(chunk_metadata)
        
        print(f"✅ Simple chunking created {len(chunks)} chunks")
        return chunks
    
    def _build_chunk_with_edges(self, chunk_blocks: List[Dict[str, Any]], all_blocks: List[Dict[str, Any]], current_index: int) -> str:
        """Build chunk content with edge windows for better context"""
        # Get the main content
        main_content = '\n\n'.join(block['content'] for block in chunk_blocks)
        
        # Add edge windows (context from surrounding blocks)
        edge_content = self._get_edge_context(chunk_blocks, all_blocks, current_index)
        
        if edge_content:
            return f"{edge_content}\n\n{main_content}"
        
        return main_content
    
    def _get_edge_context(self, chunk_blocks: List[Dict[str, Any]], all_blocks: List[Dict[str, Any]], current_index: int) -> str:
        """Get edge context from surrounding blocks"""
        edge_context = []
        edge_token_budget = self.config.CHUNK_OVERLAP // 2  # Use half overlap for edges
        
        # Get context from previous blocks (if not already included)
        if chunk_blocks and chunk_blocks[0]['position'] > 0:
            prev_block_idx = chunk_blocks[0]['position'] - 1
            if prev_block_idx >= 0 and prev_block_idx < len(all_blocks):
                prev_block = all_blocks[prev_block_idx]
                if prev_block['token_count'] <= edge_token_budget:
                    edge_context.append(f"[Context: {prev_block['content'][:200]}...]")
        
        return '\n'.join(edge_context)
    
    def _get_edge_window_blocks(self, current_blocks: List[Dict[str, Any]], all_blocks: List[Dict[str, Any]], current_index: int) -> List[Dict[str, Any]]:
        """Get blocks for edge window overlap"""
        overlap_blocks = []
        overlap_tokens = 0
        target_overlap = self.config.CHUNK_OVERLAP
        
        # Take the most important blocks from the end of current chunk
        sorted_blocks = sorted(current_blocks, key=lambda x: x['importance_score'], reverse=True)
        
        for block in sorted_blocks:
            if overlap_tokens + block['token_count'] <= target_overlap:
                overlap_blocks.append(block)
                overlap_tokens += block['token_count']
            else:
                break
        
        # Sort back by position to maintain document order
        overlap_blocks.sort(key=lambda x: x['position'])
        
        return overlap_blocks
    
    def _create_enhanced_chunk_metadata(self, content: str, source: str, chunk_id: int, blocks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Create enhanced chunk metadata with semantic information"""
        # Get block types and importance scores
        block_types = [block['block_type'] for block in blocks]
        avg_importance = sum(block['importance_score'] for block in blocks) / len(blocks) if blocks else 1.0
        
        # Create normalized content
        normalized_content = self._create_normalized_content(content)
        
        return {
            'content': content,
            'normalized_content': normalized_content,
            'source': source,
            'chunk_id': chunk_id,
            'token_count': len(self.encoding.encode(content)),
            'char_count': len(content),
            'keywords': self._extract_keywords(content),
            'block_types': block_types,
            'importance_score': avg_importance,
            'semantic_blocks_count': len(blocks)
        }
    
    def _get_overlap_sentences(self, sentence_buffer: List[str], all_sentences: List[str], current_index: int) -> List[str]:
        """Get overlap sentences that maintain semantic context"""
        # Take last few sentences as overlap, but respect semantic boundaries
        overlap_target_tokens = self.config.CHUNK_OVERLAP
        overlap_sentences = []
        total_tokens = 0
        
        # Start from the end of current buffer and work backwards
        for sentence in reversed(sentence_buffer[-5:]):  # Look at last 5 sentences max
            sentence_tokens = len(self.encoding.encode(sentence))
            if total_tokens + sentence_tokens <= overlap_target_tokens:
                overlap_sentences.insert(0, sentence)
                total_tokens += sentence_tokens
            else:
                break
        
        return overlap_sentences
    
    def _split_into_sentences_cached(self, text: str) -> List[str]:
        """Split text into sentences with caching - OPTIMIZED"""
        # Check cache first
        if text in self._token_cache:
            # Use a simple heuristic for cached results
            return [text]  # For performance, return as single sentence if cached
        
        # Use compiled regex for better performance
        sentences = self._compiled_patterns['sentence_split'].split(text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences - LEGACY METHOD"""
        return self._split_into_sentences_cached(text)
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text for chunk continuity"""
        tokens = self.encoding.encode(text)
        overlap_tokens = min(self.config.CHUNK_OVERLAP, len(tokens))
        
        if overlap_tokens > 0:
            overlap_token_ids = tokens[-overlap_tokens:]
            return self.encoding.decode(overlap_token_ids)
        
        return ""
    
    def _create_chunk_metadata(self, content: str, source: str, chunk_id: int) -> Dict[str, Any]:
        """Create chunk with enhanced metadata for better search"""
        # Create normalized version for embedding
        normalized_content = self._create_normalized_content(content)
        
        return {
            'content': content,  # Original content for display
            'normalized_content': normalized_content,  # Normalized for embedding
            'source': source,
            'chunk_id': chunk_id,
            'token_count': len(self.encoding.encode(content)),
            'char_count': len(content),
            'keywords': self._extract_keywords(content)
        }
    
    def _create_normalized_content(self, content: str) -> str:
        """Create normalized version of content for better embedding"""
        # Convert to lowercase for embedding
        normalized = content.lower()
        
        # Additional normalizations for embedding
        normalized = re.sub(r'\b(\d+)\s*(months?|years?|days?)\b', r'\1 \2', normalized)
        normalized = re.sub(r'\b(\d+)\s*%\b', r'\1 percent', normalized)
        normalized = re.sub(r'₹\s*(\d+)', r'rupees \1', normalized)
        
        return normalized
    
    def _extract_keywords(self, content: str) -> List[str]:
        """Extract important keywords from content for enhanced search"""
        # Define important terms that should be preserved
        important_terms = [
            'grace period', 'waiting period', 'pre-existing', 'coverage', 'benefit',
            'premium', 'claim', 'deductible', 'copayment', 'maternity', 'dental',
            'surgery', 'hospitalization', 'outpatient', 'inpatient', 'ayush',
            'no claim discount', 'ncd', 'icu', 'room rent', 'sub-limit'
        ]
        
        keywords = []
        content_lower = content.lower()
        
        for term in important_terms:
            if term in content_lower:
                keywords.append(term)
        
        # Extract numbers with units (periods, amounts, percentages)
        number_patterns = [
            r'\d+\s*(months?|years?|days?)',
            r'\d+\s*%',
            r'₹\s*\d+[,\d]*',
            r'\d+\s*(lakh|crore)'
        ]
        
        for pattern in number_patterns:
            matches = re.findall(pattern, content_lower)
            keywords.extend(matches)
        
        return list(set(keywords))  # Remove duplicates
    
    def clear_cache(self):
        """Clear tokenization cache to prevent memory issues"""
        self._token_cache.clear()
    
    def get_cache_stats(self) -> Dict[str, int]:
        """Get cache statistics for monitoring"""
        return {
            'cache_size': len(self._token_cache),
            'cache_memory_estimate': sum(len(k) + 4 for k in self._token_cache.keys())  # Rough estimate
        }
